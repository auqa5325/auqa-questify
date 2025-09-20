import os
import time
import json
import math
import pathlib
import random
import io
import textwrap
from docx.shared import Cm

import boto3
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from opensearchpy import OpenSearch, RequestsHttpConnection
from requests_aws4auth import AWS4Auth
from bedrockModels import build_request, count_tokens
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown2

# -----------------------------------------------------------------------------
# AUQA Question Paper Generator (v3) - Up to Step 4
# -----------------------------------------------------------------------------

# ---------------- Helpers ----------------------------------------------------

def clean_json_output(text: str) -> str:
    """Strip fence markers like ```json ... ``` or ``` ... ``` from model outputs."""
    if not isinstance(text, str):
        return text
    t = text.strip()
    if t.startswith("```json"):
        t = t[len("```json"):]
    if t.startswith("```"):
        t = t[3:]
    if t.endswith("```"):
        t = t[:-3]
    return t.strip()


def truncate_to_limit(text: str, max_tokens: int, buffer: int = 2500):
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        tokens = enc.encode(text)
        if len(tokens) > (max_tokens - buffer):
            return enc.decode(tokens[: max_tokens - buffer]), True
        return text, False
    except Exception:
        return text, False


def normalize_ratios(e: float, m: float, h: float):
    s = e + m + h
    if s <= 0:
        return 1/3, 1/3, 1/3
    return e/s, m/s, h/s


def integer_distribute_equal(total: int, groups: int):
    """Distribute `total` items as evenly as possible across `groups` buckets."""
    if groups <= 0:
        return []
    base = total // groups
    rem = total % groups
    dist = [base + (1 if i < rem else 0) for i in range(groups)]
    return dist


# ---------------- Environment & Clients -------------------------------------
load_dotenv()

OUT_DIR = os.environ.get("AUQA_OUT_DIR", "/tmp/auqa_output")
pathlib.Path(OUT_DIR).mkdir(parents=True, exist_ok=True)

region = os.environ.get("AWS_REGION", "us-east-1")
bucket_name = os.environ.get("S3_BUCKET", "my-bucket")
os_domain = os.environ.get("OS_DOMAIN", "opensearch-domain")
index_name = os.environ.get("INDEX_NAME", "test-auqa")

session = boto3.Session(region_name=region)
credentials = session.get_credentials().get_frozen_credentials()
awsauth = AWS4Auth(
    credentials.access_key, credentials.secret_key,
    region, "es", session_token=credentials.token
)
client = OpenSearch(hosts=[{"host": os_domain, "port": 443}],
                    http_auth=awsauth,
                    use_ssl=True, verify_certs=True,
                    connection_class=RequestsHttpConnection)

s3 = boto3.client("s3", region_name=region)
textract = boto3.client("textract", region_name=region)
bedrock = session.client("bedrock-runtime", region_name=region)

# ---------------- Streamlit UI ----------------------------------------------
st.set_page_config(layout="wide")
st.title("📘 AUQA: Difficulty-aware Question Paper Generator (v3)")

subject_code = st.text_input("Subject Code", "CS23304")
subject_name = st.text_input("Subject Name", "Java Programming")
department = st.text_input("Department", "Computer Science and Engineering")
semester = st.text_input("Semester", "III")
regulation = st.text_input("Regulation", "2023")
exam_year = st.text_input("Exam Session (e.g. NOV/DEC 2025)", "NOV/DEC 2025")

s3_key = st.text_input("S3 PDF Key (syllabus)", "syllabus/CS23501.Syllabus.pdf")

st.markdown("### Overall Difficulty Ratio (applies to PART A and PART B)")
col1, col2, col3 = st.columns(3)
_easy = col1.number_input("Easy ratio", 0.0, 1.0, 0.34, 0.01)
_medium = col2.number_input("Medium ratio", 0.0, 1.0, 0.33, 0.01)
_hard = col3.number_input("Hard ratio", 0.0, 1.0, 0.33, 0.01)
_e, _m, _h = normalize_ratios(_easy, _medium, _hard)

# Models config
BASE_DIR = pathlib.Path(__file__).resolve().parent.parent
MODEL_FILE = BASE_DIR / "models.json"
if MODEL_FILE.exists():
    with open(MODEL_FILE, "r") as f:
        MODELS = json.load(f)
    model_names = [m["name"] for m in MODELS]
    selected_model = st.selectbox("Choose Model:", model_names)
    model_config = next(m for m in MODELS if m["name"] == selected_model)
else:
    st.warning("models.json not found. LLM invocation will be disabled.")
    MODELS, model_config = [], None

if "units_parsed" not in st.session_state:
    st.session_state.units_parsed = []
if "COs" not in st.session_state:
    st.session_state.COs = []
if "final_paper" not in st.session_state:
    st.session_state.final_paper = []

# ---------------- Step 1: Ingest syllabus -----------------------------------
if st.button("Ingest syllabus PDF from S3 and parse units & COs"):
    st.info(f"Starting Textract job for s3://{bucket_name}/{s3_key}")
    try:
        resp = textract.start_document_text_detection(
            DocumentLocation={"S3Object": {"Bucket": bucket_name, "Name": s3_key}}
        )
        job_id = resp["JobId"]
        while True:
            status_resp = textract.get_document_text_detection(JobId=job_id)
            status = status_resp.get("JobStatus")
            if status in ("SUCCEEDED", "FAILED"):
                break
            time.sleep(3)

        if status == "FAILED":
            st.error("Textract job failed")
        else:
            page_texts = {}
            next_token = None
            while True:
                if next_token:
                    chunk = textract.get_document_text_detection(JobId=job_id, NextToken=next_token)
                else:
                    chunk = textract.get_document_text_detection(JobId=job_id)
                for block in chunk.get("Blocks", []):
                    if block.get("BlockType") == "LINE":
                        page_no = block.get("Page")
                        page_texts.setdefault(page_no, []).append(block.get("Text"))
                next_token = chunk.get("NextToken")
                if not next_token:
                    break
            full_text = "\n\n".join("\n".join(page_texts[p]) for p in sorted(page_texts.keys()))

            # Parse prompt template
            parse_template = """
INSTRUCTION:
You are an assistant specialized in extracting a course outline.

INPUT:
- Course ID: {SUBJECT_CODE}
- Document: extracted text below.

TASK:
Return ONLY a JSON object:
{
  "course_id": "{SUBJECT_CODE}",
  "course_objectives": [
     { "id": "CO1", "description": "..." }
  ],
  "units": [
    { "unit_no": <int>, "unit_name": "<string>", "topics": ["<string>", ...] }
  ]
}

RULES:
- Output JSON only (no commentary).
- Integers for unit_no.
- If missing, output empty lists.
- Ensure valid JSON.

---DOCUMENT-BEGIN---
{FULL_TEXT}
---DOCUMENT-END---
"""
            parse_prompt = parse_template.replace("{SUBJECT_CODE}", subject_code)\
                                         .replace("{FULL_TEXT}", full_text)

            if model_config:
                body = build_request(model_config["id"], selected_model, parse_prompt, 4096)
                resp = bedrock.invoke_model(modelId=model_config["id"], body=json.dumps(body))
                model_response = json.loads(resp["body"].read())
                generated_text = (model_response.get("outputs", [{}])[0].get("text") or
                                  model_response.get("content", [{}])[0].get("text") or
                                  str(model_response))
                generated_text = clean_json_output(generated_text)
                st.subheader("Parsed Units + COs JSON (raw)")
                st.code(generated_text)
                try:
                    parsed = json.loads(generated_text)
                    st.session_state.units_parsed = parsed.get("units", [])
                    st.session_state.COs = parsed.get("course_objectives", [])
                    st.success("Parsed syllabus successfully")
                except Exception as e:
                    st.error(f"Failed to parse JSON: {e}")
                    st.text_area("Model output", generated_text, height=300)
            else:
                st.warning("Model config missing; saving raw text only.")
                st.session_state.raw_text = full_text
    except Exception as e:
        st.error(f"Textract error: {e}")

# ---------------- Step 2: Select Units -------------------------------------
if st.session_state.units_parsed:
    st.subheader("Select Units for Question Generation")
    selected_units = []
    for unit in st.session_state.units_parsed:
        unit_no = unit.get("unit_no")
        unit_name = unit.get("unit_name")
        with st.expander(f"Unit {unit_no}: {unit_name}"):
            if st.checkbox("Include this unit", key=f"enable_{unit_no}"):
                selected_units.append(unit)
    st.write(f"Selected units: {[u['unit_no'] for u in selected_units]}")

# ---------------- Step 3: Retrieve chunks ----------------------------------
if st.button("Retrieve contexts for selected units") and st.session_state.units_parsed:
    if not selected_units:
        st.warning("No units selected")
    else:
        unit_contexts = {}
        for unit in selected_units:
            unit_no = unit["unit_no"]
            unit_name = unit["unit_name"]
            query = f"{subject_code} " + (" ".join(unit.get("topics", [])) or unit_name)

            # BM25
            try:
                bm25_query = {"size": 5, "query": {"match": {"chunk_text": query}}}
                bm25_resp = client.search(index=index_name, body=bm25_query)
                bm25_docs = [(hit["_id"], hit["_source"]["chunk_text"], hit["_source"].get("page_range", "?"))
                             for hit in bm25_resp["hits"]["hits"]]
            except Exception as e:
                st.warning(f"BM25 search failed for unit {unit_no}: {e}")
                bm25_docs = []

            # Vector
            try:
                embedding_body = json.dumps({"inputText": query})
                resp = bedrock.invoke_model(modelId="amazon.titan-embed-text-v2:0", body=embedding_body)
                emb = json.loads(resp["body"].read())["embedding"]
                vector_query = {"size": 5, "query": {"knn": {"vector_field": {"vector": emb, "k": 5}}}}
                vector_resp = client.search(index=index_name, body=vector_query)
                vector_docs = [(hit["_id"], hit["_source"]["chunk_text"], hit["_source"].get("page_range", "?"))
                               for hit in vector_resp.get("hits", {}).get("hits", [])]
            except Exception as e:
                st.warning(f"Vector search failed for unit {unit_no}: {e}")
                vector_docs = []

            selected = bm25_docs[:1] + vector_docs[:4]
            unit_contexts[unit_no] = {"unit_name": unit_name, "selected_chunks": selected}
            if selected:
                df = pd.DataFrame(selected, columns=["Doc ID", "Chunk Text", "Page Range"])
                st.subheader(f"Unit {unit_no} chunks")
                st.dataframe(df)
            else:
                st.warning(f"No chunks found for unit {unit_no}")
        st.session_state.unit_contexts = unit_contexts

# ---------------- Step 4: Prompt Templates ---------------------------------
# ---------------- Step 4: Build and send prompts to LLM ---------------------

PART_A_SLOTS = 10
PART_B_QUESTIONS = 5
PART_B_SUBPARTS = PART_B_QUESTIONS * 2
PART_C_SUBPARTS = 2

# Prompt templates
part_a_template = """
You are an AI assistant specialized in exam question generation.
Generate {N_A} short questions (2 marks each) for PART A for Unit {UNIT_NO}: {UNIT_NAME}.
Guidelines:
- SHORT to write, reasoning provides difficulty.
- Difficulty ratio across paper: Easy={E}, Medium={M}, Hard={H}.
- Map each to a CO where possible.
Return JSON array: question_text, difficulty, blooms_level, CO, marks, part='A', q_no, subpart=null.
"""

part_b_template = """
You are an AI assistant specialized in exam question generation.
Generate {N_B} subpart questions (13 marks each) for PART B for Unit {UNIT_NO}: {UNIT_NAME}.
Guidelines:
- May/may not have subdivisions (a,b).
- Moderate time to solve.
- Difficulty = reasoning effort, not length.
Return JSON array: question_text, difficulty, blooms_level, CO, marks=13, part='B', q_no, subpart=null.
"""

part_c_template = """
You are an AI assistant specialized in exam question generation.
Generate 1 HARD question for PART C subpart {SUB_ID} (marks={MARKS}) based on Unit {UNIT_NO}: {UNIT_NAME}.
Guidelines:
- Scenario/problem-based requiring deduction/application.
- Strictly HARD difficulty.
Return JSON object: question_text, difficulty='Hard', blooms_level, CO, marks={MARKS}, part='C', q_no=16, subpart='{SUB_ID}'.
"""

if "unit_contexts" in st.session_state and st.session_state.unit_contexts:
    if st.button("🚀 Generate Full Question Paper"):
        unit_contexts = st.session_state.unit_contexts
        selected_unit_nos = list(unit_contexts.keys())
        n_units = len(selected_unit_nos)
        if n_units == 0:
            st.error("No units selected for generation")
        else:
            # Distribute Part A slots equally across selected units
            part_a_distribution = integer_distribute_equal(PART_A_SLOTS, n_units)
            # Distribute Part B subparts equally
            part_b_distribution = integer_distribute_equal(PART_B_SUBPARTS, n_units)

            st.write("Part A distribution per unit:", dict(zip(selected_unit_nos, part_a_distribution)))
            st.write("Part B distribution per unit (subparts):", dict(zip(selected_unit_nos, part_b_distribution)))

            all_questions = []
            cos = st.session_state.get("COs", [])
            cos_text = "\n".join([f"{c.get('id')}: {c.get('description')}" for c in cos]) if cos else ""

            # ---------------- Part A ----------------
            part_a_counter = 1
            for idx, unit_no in enumerate(selected_unit_nos):
                ctx = unit_contexts[unit_no]
                unit_name = ctx["unit_name"]
                chunks = ctx["selected_chunks"]
                context_text = "\n\n".join([t for _, t, _ in chunks])
                context_refs = [p for _, _, p in chunks]

                n_a = part_a_distribution[idx]
                if n_a > 0:
                    llm_guidance = f"""
You are an AI assistant specialized in exam question generation.
Generate {n_a} short questions (2 marks each) for PART A for Unit {unit_no}: {unit_name}.
Guidelines:
- Questions must be short to write, but difficulty should reflect reasoning required.
- Use difficulty ratio across whole paper: Easy={_e:.2f}, Medium={_m:.2f}, Hard={_h:.2f}.
- Map questions to COs if possible.
Return ONLY a JSON array of objects:
[{{"question_text":"...", "difficulty":"...", "blooms_level":"...", "CO":"...", "marks":2,
   "part":"A", "q_no":<int>, "subpart":null, "context_ref":"..."}}]
"""
                    prompt_obj = {
                        "llm_guidance": llm_guidance,
                        "paper_context": {
                            "unit_no": unit_no,
                            "unit_name": unit_name,
                            "context": context_text,
                            "pages": context_refs,
                            "course_objectives": cos
                        }
                    }
                    full_prompt = json.dumps(prompt_obj, indent=2)
                    full_prompt, _ = truncate_to_limit(full_prompt,
                                                       model_config.get("max_tokens", 4096) if model_config else 2048)
                    if model_config:
                        body = build_request(model_config["id"], selected_model, full_prompt, 4096)
                        resp = bedrock.invoke_model(modelId=model_config["id"], body=json.dumps(body))
                        model_response = json.loads(resp["body"].read())
                        gtext = (model_response.get("outputs", [{}])[0].get("text")
                                 or model_response.get("content", [{}])[0].get("text")
                                 or str(model_response))
                        gtext = clean_json_output(gtext)
                        try:
                            items = json.loads(gtext)
                            for it in items:
                                it.setdefault("marks", 2)
                                it.setdefault("part", "A")
                                it["q_no"] = part_a_counter
                                part_a_counter += 1
                                it.setdefault("subpart", None)
                                all_questions.append(it)
                            st.success(f"Generated {len(items)} Part-A questions for Unit {unit_no}")
                        except Exception as e:
                            st.error(f"Failed to parse Part-A JSON for unit {unit_no}: {e}")
                            st.text_area("LLM output", gtext, height=200)

            # ---------------- Part B ----------------
            part_b_counter = 11
            b_subparts = []
            for idx, unit_no in enumerate(selected_unit_nos):
                ctx = unit_contexts[unit_no]
                unit_name = ctx["unit_name"]
                chunks = ctx["selected_chunks"]
                context_text = "\n\n".join([t for _, t, _ in chunks])
                context_refs = [p for _, _, p in chunks]

                n_b_subparts = part_b_distribution[idx]
                if n_b_subparts > 0:
                    llm_guidance = f"""
You are an AI assistant specialized in exam question generation.
Generate {n_b_subparts} subpart questions (13 marks each) for PART B for Unit {unit_no}: {unit_name}.
Guidelines:
- These will later be grouped as (a)/(b).
- Questions should be moderate in depth (longer than Part A, shorter than Part C).
- Use difficulty ratio across whole paper: Easy={_e:.2f}, Medium={_m:.2f}, Hard={_h:.2f}.
- Map to COs where possible.
Return ONLY a JSON array of objects:
[{{"question_text":"...", "difficulty":"...", "blooms_level":"...", "CO":"...", "marks":13,
   "part":"B", "q_no":<int>, "subpart":null}}]
"""
                    prompt_obj = {
                        "llm_guidance": llm_guidance,
                        "paper_context": {
                            "unit_no": unit_no,
                            "unit_name": unit_name,
                            "context": context_text,
                            "pages": context_refs,
                            "course_objectives": cos
                        }
                    }
                    full_prompt = json.dumps(prompt_obj, indent=2)
                    full_prompt, _ = truncate_to_limit(full_prompt,
                                                       model_config.get("max_tokens", 4096) if model_config else 2048)
                    if model_config:
                        body = build_request(model_config["id"], selected_model, full_prompt, 4096)
                        resp = bedrock.invoke_model(modelId=model_config["id"], body=json.dumps(body))
                        model_response = json.loads(resp["body"].read())
                        gtext = (model_response.get("outputs", [{}])[0].get("text")
                                 or model_response.get("content", [{}])[0].get("text")
                                 or str(model_response))
                        gtext = clean_json_output(gtext)
                        try:
                            items = json.loads(gtext)
                            for it in items:
                                it.setdefault("marks", 13)
                                it.setdefault("part", "B")
                                it.setdefault("q_no", 0)  # temporary, will fix after grouping
                                it.setdefault("subpart", None)
                                b_subparts.append(it)
                            st.success(f"Generated {len(items)} Part-B subparts for Unit {unit_no}")
                        except Exception as e:
                            st.error(f"Failed to parse Part-B JSON for unit {unit_no}: {e}")
                            st.text_area("LLM output", gtext, height=200)

            # Group Part B into Q11–Q15 with (a)/(b)
            grouped_B = []
            for i in range(0, len(b_subparts), 2):
                a = b_subparts[i]
                b = b_subparts[i + 1] if i + 1 < len(b_subparts) else None
                a["q_no"] = part_b_counter
                a["subpart"] = "a"
                if b:
                    b["q_no"] = part_b_counter
                    b["subpart"] = "b"
                grouped_B.append((a, b))
                part_b_counter += 1

            for pair in grouped_B:
                all_questions.extend([q for q in pair if q])

            # ---------------- Part C ----------------
            part_c_items = []
            for sub_id, marks in zip(["a", "b"], [8, 7]):
                chosen_unit = random.choice(selected_unit_nos)
                ctx = unit_contexts[chosen_unit]
                unit_name = ctx["unit_name"]
                chunks = ctx["selected_chunks"]
                context_text = "\n\n".join([t for _, t, _ in chunks])
                context_refs = [p for _, _, p in chunks]

                llm_guidance = f"""
You are an AI assistant specialized in exam question generation.
Generate 1 HARD question for PART C subpart {sub_id} (marks={marks}) based on Unit {chosen_unit}: {unit_name}.
Guidelines:
- Scenario or problem-based, requiring deep reasoning.
- Strictly HARD difficulty.
Return ONLY a JSON object:
{{"question_text":"...", "difficulty":"Hard", "blooms_level":"...", "CO":"...", "marks":{marks},
  "part":"C", "q_no":16, "subpart":"{sub_id}", "context_ref":"..."}}
"""
                prompt_obj = {
                    "llm_guidance": llm_guidance,
                    "paper_context": {
                        "unit_no": chosen_unit,
                        "unit_name": unit_name,
                        "context": context_text,
                        "pages": context_refs,
                        "course_objectives": cos
                    }
                }
                full_prompt = json.dumps(prompt_obj, indent=2)
                full_prompt, _ = truncate_to_limit(full_prompt,
                                                   model_config.get("max_tokens", 4096) if model_config else 2048)
                if model_config:
                    body = build_request(model_config["id"], selected_model, full_prompt, 4096)
                    resp = bedrock.invoke_model(modelId=model_config["id"], body=json.dumps(body))
                    model_response = json.loads(resp["body"].read())
                    gtext = (model_response.get("outputs", [{}])[0].get("text")
                             or model_response.get("content", [{}])[0].get("text")
                             or str(model_response))
                    gtext = clean_json_output(gtext)
                    try:
                        obj = json.loads(gtext)
                        obj.setdefault("marks", marks)
                        obj.setdefault("part", "C")
                        obj.setdefault("q_no", 16)
                        obj.setdefault("subpart", sub_id)
                        part_c_items.append(obj)
                        st.success(f"Generated Part-C {sub_id} from unit {chosen_unit}")
                    except Exception as e:
                        st.error(f"Failed to parse Part-C JSON for subpart {sub_id}: {e}")
                        st.text_area("LLM output", gtext, height=200)

            all_questions.extend(part_c_items)

            st.header("Generated Question Paper JSON")
            st.json(all_questions)

            # Save to S3
            out_key = f"generated/{subject_code}_{exam_year.replace('/','_')}_questions.json"
            

            st.session_state.final_paper = all_questions


# ---------------- Step 5: Export / Display Generated Paper ------------------
def export_final_paper(final_paper_list: list,
                       out_dir_path: str = None,
                       subject_name_local: str = None,
                       subject_code_local: str = None,
                       semester_local: str = None):
    """Export the generated paper as CSV, MD, and DOCX with simple line-by-line questions (no tables)."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    out_dir = pathlib.Path(out_dir_path or OUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)

    if not final_paper_list:
        raise ValueError("final_paper_list is empty")

    # Save CSV
    df = pd.DataFrame(final_paper_list)
    csv_path = out_dir / "generated_questions.csv"
    df.to_csv(csv_path, index=False)

    # Save MD
    subj = subject_name_local or subject_name
    scode = subject_code_local or subject_code
    sem = semester_local or semester
    md_text = f"# {subj} - Generated Question Paper\n\n**Subject Code:** {scode}  \n**Semester:** {sem}\n"
    md_path = out_dir / "generated_paper.md"
    md_path.write_text(md_text, encoding="utf-8")

    # DOCX setup
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Header
    p = doc.add_paragraph("ANNA UNIVERSITY (UNIVERSITY DEPARTMENTS)")
    p.runs[0].bold, p.runs[0].font.size, p.alignment = True, Pt(14), WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(f"B.E. (Full Time) - END SEMESTER EXAMINATIONS, {exam_year}")
    p.runs[0].font.size, p.alignment = Pt(12), WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph(f"{subj} - {department} - Semester: {sem}  Reg: {regulation}")
    p.runs[0].bold, p.alignment = True, WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("Time: 3 Hrs\t\tMax.Marks: 100")

    # Course Objectives
    cos = st.session_state.get("COs", [])
    if cos:
        p = doc.add_paragraph("Course Objectives:")
        p.runs[0].bold = True
        for co in cos:
            doc.add_paragraph(f"{co.get('id', '')}. {co.get('description', '')}")

    def format_question(item, qno_label):
        co_val = item.get("CO", "-")
        bl_val = item.get("blooms_level", "-")
        marks_val = item.get("marks", "-")
        return f"{qno_label}. {item.get('question_text', '')} (CO: {co_val}, BL: {bl_val}, Marks: {marks_val})"

    # Part A
    doc.add_paragraph("PART - A (10 x 2 = 20 Marks)").runs[0].bold = True
    doc.add_paragraph("(Answer all Questions)")
    a_items = sorted([q for q in final_paper_list if q.get("part") == "A"], key=lambda x: x.get("q_no", 0))[:10]
    for it in a_items:
        doc.add_paragraph(format_question(it, it.get("q_no")))

    # Part B
    doc.add_paragraph("PART - B (5 x 13 = 65 Marks)").runs[0].bold = True
    doc.add_paragraph("(Restrict to a maximum of 2 subdivisions)")
    b_groups = {}
    for q in [q for q in final_paper_list if q.get("part") == "B"]:
        b_groups.setdefault(int(q.get("q_no", 0)), []).append(q)
    for qno in sorted(b_groups.keys()):
        for p in sorted(b_groups[qno], key=lambda x: x.get("subpart") or "a"):
            label = f"{qno}({p.get('subpart')})"
            doc.add_paragraph(format_question(p, label))

    # Part C
    doc.add_paragraph("PART - C (1 x 15 = 15 Marks)").runs[0].bold = True
    doc.add_paragraph("(Q.No.16 is compulsory)")
    c_items = sorted([q for q in final_paper_list if q.get("part") == "C"], key=lambda x: x.get("subpart") or "a")
    for p in c_items:
        label = f"16({p.get('subpart')})"
        doc.add_paragraph(format_question(p, label))

    doc.add_paragraph("Note: This question paper was auto-generated by AUQA.")

    docx_path = out_dir / f"{subject_code}_{exam_year.replace('/', '_')}_paper.docx"
    doc.save(str(docx_path))

    try:
        html = markdown2.markdown(md_text)
        html_path = out_dir / "generated_paper.html"
        html_path.write_text(html, encoding="utf-8")
    except Exception:
        html_path = None

    return {"csv": str(csv_path), "md": str(md_path), "docx": str(docx_path), "html": str(html_path) if html_path else None}

# Export / Download UI
if "final_paper" in st.session_state and st.session_state.final_paper:
    st.markdown("---")
    st.subheader("Export / Download")
    if st.button("Prepare export (CSV/MD + DOCX download)"):
        try:
            results = export_final_paper(st.session_state.final_paper)
            st.success("Prepared export. You can download the DOCX directly below (not saved permanently).")
            # show CSV/MD paths
            st.write({"csv": results["csv"], "md": results["md"], "docx": results["docx"], "html": results["html"]})
            # provide download button for DOCX bytes
            with open(results["docx"], "rb") as f:
                docx_bytes = f.read()
            st.download_button(label="Download DOCX", data=docx_bytes,
                               file_name=f"{subject_code}_{exam_year.replace('/','_')}_paper.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"Export failed: {e}")

# End of file
